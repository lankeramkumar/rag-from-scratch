"""
Author: Ram Kumar Lanke
Applied AI Solutions Architect
lankeramkumar@gmail.com

Generates the 20 sample source documents used throughout this RAG mini-course.

Run once:
    python scripts/00_generate_sample_documents.py

Produces documents/ containing a deliberately mixed bag of file types
(.txt, .pdf, .png, .py/.js/.sql/.sh, .docx) so the text-extraction step
(01_text_extraction.py) has something realistic to normalize.
"""

import os
import textwrap

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(HERE, "..", "documents")
os.makedirs(DOCS_DIR, exist_ok=True)


def path(name: str) -> str:
    return os.path.join(DOCS_DIR, name)


# ---------------------------------------------------------------------------
# 1. Plain text documents (5)
# ---------------------------------------------------------------------------

TEXT_DOCS = {
    "doc01_solar_system.txt": """The Solar System: An Overview

The Solar System formed roughly 4.6 billion years ago from the gravitational
collapse of a giant interstellar molecular cloud. Most of the mass collected
at the center, forming the Sun, while the remaining material flattened into
a protoplanetary disk out of which the planets, moons, asteroids, and comets
eventually formed.

The Sun accounts for about 99.86% of the Solar System's total mass. It is a
G-type main-sequence star that fuses roughly 600 million tons of hydrogen
into helium every second, releasing the energy that powers almost every
process on Earth, from weather to photosynthesis.

There are eight recognized planets. The four inner, or terrestrial, planets
-- Mercury, Venus, Earth, and Mars -- are small, dense, and rocky. Mercury is
the closest to the Sun and has almost no atmosphere, leading to extreme
temperature swings between its day and night sides. Venus, similar in size
to Earth, has a runaway greenhouse atmosphere of carbon dioxide that makes
its surface hot enough to melt lead. Earth is the only known planet to
harbor life, protected by a magnetic field and a nitrogen-oxygen atmosphere.
Mars, the "Red Planet," shows strong evidence of ancient rivers and lakes
and is the most actively explored planet besides Earth, with several rovers
including Curiosity and Perseverance studying its geology for signs of past
microbial life.

Beyond the asteroid belt lie the four giant planets. Jupiter is the largest
planet in the Solar System, more massive than all the other planets
combined. Its Great Red Spot is a storm that has raged for at least 350
years. Saturn is famous for its spectacular ring system, made mostly of
ice particles with a smaller amount of rocky debris and dust. Uranus and
Neptune are classified as ice giants; Uranus is unusual in that it rotates
almost on its side, likely the result of an ancient collision.

The asteroid belt between Mars and Jupiter contains millions of rocky
bodies, the largest of which, Ceres, is classified as a dwarf planet.
Beyond Neptune lies the Kuiper Belt, a disc-shaped region containing many
small icy bodies, including the dwarf planet Pluto, which was reclassified
from a full planet in 2006 after the discovery of other similarly sized
Kuiper Belt objects.

Comets are icy bodies that, when they approach the Sun, develop glowing
comas and often long tails of gas and dust blown by the solar wind. Halley's
Comet, the most famous, returns to the inner Solar System roughly every 76
years.

Moons orbit most of the planets. Earth has one large moon, thought to have
formed from debris after a Mars-sized body collided with the young Earth.
Jupiter has over 90 known moons, including the four large Galilean moons
discovered by Galileo Galilei in 1610: Io, Europa, Ganymede, and Callisto.
Europa is of particular scientific interest because it likely harbors a
liquid water ocean beneath its icy crust, making it one of the leading
candidates in the search for extraterrestrial life.

Space exploration has transformed our understanding of the Solar System.
The Voyager 1 and 2 probes, launched in 1977, have both left the
heliosphere and entered interstellar space, still transmitting data after
more than four decades. Modern missions continue to refine our knowledge:
the Parker Solar Probe has flown closer to the Sun than any spacecraft in
history, and the James Webb Space Telescope, while primarily focused on
deep space, has also captured striking new images of planets within our
own Solar System.

Understanding the Solar System is not just an academic exercise. Near-Earth
objects pose a genuine, if statistically small, risk of impact, and
planetary science informs our models of climate, atmospheric chemistry, and
the conditions necessary for life -- questions with direct relevance back
here on Earth.
""",
    "doc02_python_programming.txt": """An Introduction to the Python Programming Language

Python is a high-level, general-purpose programming language created by
Guido van Rossum and first released in 1991. Its design philosophy
emphasizes code readability, using significant indentation to delimit code
blocks rather than curly braces or keywords. This makes Python one of the
most approachable languages for beginners while remaining powerful enough
for large-scale production systems.

Python is dynamically typed and garbage-collected. It supports multiple
programming paradigms, including procedural, object-oriented, and
functional programming. Variables do not need explicit type declarations;
the interpreter infers types at runtime. This flexibility speeds up
development but places more responsibility on testing to catch type-related
errors that a statically typed language would catch at compile time.

One of Python's greatest strengths is its extensive standard library,
often described with the phrase "batteries included." Modules for file
I/O, networking, regular expressions, data serialization (json, csv,
pickle), and concurrency are all available without installing anything
extra. Beyond the standard library, the Python Package Index (PyPI) hosts
hundreds of thousands of third-party packages, installable via pip.

Python has become the dominant language in data science and machine
learning. Libraries such as NumPy provide efficient array operations
backed by compiled C code, pandas offers powerful data manipulation
through DataFrame objects, and frameworks like scikit-learn, PyTorch, and
TensorFlow make it straightforward to build everything from a simple
linear regression to a large neural network. This ecosystem is one of the
main reasons Python is a natural fit for building Retrieval-Augmented
Generation (RAG) pipelines: the same language can handle document parsing,
numerical embedding, vector search, and orchestration of large language
model calls.

Web development is another major use case. Frameworks like Django provide
a full-featured, "batteries included" approach with a built-in ORM,
authentication system, and admin interface, while Flask and FastAPI offer
lighter-weight, more flexible alternatives. FastAPI in particular has
grown rapidly in popularity because of its native support for asynchronous
request handling and automatic OpenAPI documentation generation, which
pairs well with building APIs around LLM applications.

Python's automation and scripting capabilities make it a favorite for
DevOps and system administration tasks. Tools like Ansible are written in
Python, and countless organizations use Python scripts to glue together
disparate systems, process log files, or orchestrate deployment pipelines.

Performance is often cited as a weakness of Python, since the reference
implementation (CPython) is an interpreted language and includes a Global
Interpreter Lock (GIL) that limits true multi-threaded parallelism for
CPU-bound code. In practice, this is mitigated in several ways: I/O-bound
workloads benefit from Python's asyncio library and cooperative
multitasking; CPU-bound numerical work is typically delegated to
compiled libraries like NumPy; and for problems that need genuine
parallel execution, the multiprocessing module spins up separate
processes, each with its own interpreter and memory space, sidestepping
the GIL entirely.

Python's package and environment management has matured significantly.
Tools like venv and virtualenv isolate dependencies per project, while
newer tools such as Poetry and uv aim to unify dependency resolution,
packaging, and virtual environment management into a single, faster
workflow. Choosing a consistent environment strategy is especially
important for reproducible RAG pipelines, where a mismatched version of
a tokenizer or embedding library can silently change results.

Because of its readability, enormous ecosystem, and strong community
support, Python remains one of the most widely taught first programming
languages and, simultaneously, one of the most heavily used languages in
production AI systems today.
""",
    "doc03_french_revolution.txt": """The French Revolution: Causes, Course, and Consequences

The French Revolution, which lasted from 1789 to roughly 1799, was one of
the most significant political and social upheavals in modern history. It
overturned the French monarchy, dismantled centuries-old feudal
structures, and introduced ideas about citizenship and inalienable rights
that continue to shape political thought worldwide.

The roots of the Revolution lay in a combination of financial crisis and
social inequality. France's involvement in costly wars, including
significant financial support for the American Revolution, had left the
royal treasury nearly bankrupt. At the same time, French society was
divided into three estates: the clergy, the nobility, and the commoners
(the Third Estate), who made up roughly 97% of the population yet bore
the overwhelming majority of the tax burden. Poor harvests in the 1780s
drove up bread prices, deepening resentment among ordinary people already
struggling under a regressive tax system.

In May 1789, King Louis XVI convened the Estates-General for the first
time since 1614 to address the fiscal crisis. Disputes over voting
procedures led the Third Estate to break away and declare itself the
National Assembly in June 1789, asserting that it represented the will of
the French people. When rumors spread that the King intended to
suppress this new assembly by force, Parisians stormed the Bastille
fortress-prison on July 14, 1789 -- an event that has since become the
symbolic beginning of the Revolution and is commemorated annually as
Bastille Day.

In August 1789, the National Assembly adopted the Declaration of the
Rights of Man and of the Citizen, a foundational document asserting that
all men are born free and equal in rights, and that sovereignty resides
in the nation rather than the monarch. This period, often called the
"Liberal Revolution," saw the drafting of a constitutional monarchy,
though tensions with the King and foreign powers continued to escalate.

The Revolution radicalized sharply from 1792 onward. France was declared
a republic in September 1792, and King Louis XVI was tried for treason
and executed by guillotine in January 1793. The ensuing Reign of Terror
(1793-1794), led largely by Maximilien Robespierre and the Committee of
Public Safety, saw tens of thousands of people executed or imprisoned as
suspected enemies of the Revolution. The Terror ended with Robespierre's
own arrest and execution in July 1794 (the Thermidorian Reaction),
ushering in a more conservative, if unstable, phase of government known
as the Directory.

The instability of the Directory period created an opening for a young,
ambitious general named Napoleon Bonaparte, who had risen to prominence
through military victories in Italy and Egypt. In November 1799,
Napoleon staged a coup d'etat (the 18 Brumaire) that effectively ended
the Revolution and installed him as First Consul, setting the stage for
his eventual self-coronation as Emperor of the French in 1804.

The consequences of the French Revolution were profound and long-lasting.
It ended feudal privileges and hereditary nobility as the basis of legal
status in France, established the principle of legal equality before the
law, and spread revolutionary and nationalist ideas across Europe through
the Napoleonic Wars that followed. The metric system, secular public
education, and the Napoleonic Code -- a unified civil law system that
influenced legal codes across the world -- all trace their origins to this
turbulent decade. At the same time, the Revolution's violent excesses,
particularly during the Terror, remain a subject of ongoing historical
debate about the costs of radical political change.
""",
    "doc04_nutrition_basics.txt": """Nutrition Basics: Understanding Macronutrients and a Balanced Diet

Good nutrition is one of the most powerful and cost-effective tools for
maintaining long-term health. At its core, nutrition science is about
understanding how the foods we eat are broken down and used by the body
for energy, growth, and repair.

The three macronutrients -- carbohydrates, proteins, and fats -- provide
the body with calories, or energy. Carbohydrates are the body's preferred
and most efficient energy source. They are broken down into glucose, which
fuels the brain, muscles, and other organs. Complex carbohydrates, found
in whole grains, legumes, and vegetables, are digested more slowly than
simple sugars, providing a steadier release of energy and typically
coming packaged with fiber, vitamins, and minerals. Refined carbohydrates
and added sugars, by contrast, are digested quickly and can contribute to
rapid blood sugar spikes if consumed in excess.

Proteins are composed of amino acids, nine of which are considered
"essential" because the human body cannot synthesize them and must obtain
them from food. Protein is critical for building and repairing tissue,
producing enzymes and hormones, and supporting immune function. Animal
sources such as meat, fish, eggs, and dairy provide complete proteins
containing all essential amino acids, while most plant sources need to be
combined (for example, rice and beans) to provide a complete amino acid
profile.

Dietary fats, long unfairly maligned, are essential for absorbing
fat-soluble vitamins (A, D, E, and K), building cell membranes, and
producing hormones. Unsaturated fats, found in olive oil, nuts, seeds,
and fatty fish, are generally considered beneficial for cardiovascular
health. Saturated fats, common in red meat and butter, are recommended in
moderation, while artificial trans fats -- once common in processed and
fried foods -- have been shown to raise LDL ("bad") cholesterol and are
now banned or heavily restricted in many countries.

Beyond macronutrients, micronutrients -- vitamins and minerals -- are
required in much smaller quantities but are no less essential. Vitamin C
supports collagen synthesis and immune function; vitamin D, synthesized
in the skin upon sun exposure, is critical for calcium absorption and
bone health; iron is a core component of hemoglobin, the protein that
carries oxygen in the blood; and calcium and magnesium both play roles in
bone density, muscle function, and nerve signaling.

Dietary fiber, technically a type of carbohydrate the human body cannot
fully digest, plays an outsized role in digestive health. Soluble fiber,
found in oats and legumes, can help lower cholesterol and stabilize blood
sugar, while insoluble fiber, found in whole grains and vegetable skins,
adds bulk and promotes regular bowel movements. Fiber also feeds
beneficial gut bacteria, an area of nutrition science -- the gut
microbiome -- that has seen an explosion of research interest in the
past two decades.

Hydration is often overlooked but is fundamental to nearly every bodily
process, from regulating temperature to transporting nutrients and
flushing waste. General guidance suggests roughly two to three liters of
fluid per day for most adults, though individual needs vary based on
climate, activity level, and body size.

Rather than focusing on any single nutrient in isolation, most nutrition
experts now emphasize overall dietary patterns. Diets rich in vegetables,
fruits, whole grains, lean protein, and healthy fats -- and lower in
ultra-processed foods, added sugar, and excess sodium -- are consistently
associated with reduced risk of chronic diseases including heart disease,
type 2 diabetes, and certain cancers. Small, sustainable changes, applied
consistently over time, tend to produce far better long-term outcomes
than short-lived, restrictive diets.
""",
    "doc05_climate_change.txt": """Climate Change: Mechanisms, Evidence, and Response

Climate change refers to long-term shifts in temperatures and weather
patterns, primarily driven since the mid-20th century by human activities,
especially the burning of fossil fuels. While the Earth's climate has
varied naturally over geological timescales, the current warming trend is
proceeding at a pace and scale that the overwhelming majority of climate
scientists attribute to human influence.

The physical mechanism behind human-caused warming is the greenhouse
effect. Certain gases in the atmosphere -- primarily carbon dioxide (CO2),
methane (CH4), and nitrous oxide (N2O) -- trap heat that would otherwise
radiate from the Earth's surface back into space. This effect is natural
and necessary; without any greenhouse gases, the planet's average
temperature would be far below freezing. The problem is one of degree:
since the Industrial Revolution, atmospheric CO2 concentrations have
risen from about 280 parts per million to over 420 parts per million,
primarily due to burning coal, oil, and natural gas for energy, along
with deforestation and industrial processes like cement production.

The evidence for warming comes from multiple independent lines of data.
Global average surface temperatures have risen by roughly 1.1 to 1.2
degrees Celsius since the late 19th century. Ice cores drilled from
Antarctica and Greenland preserve tiny bubbles of ancient atmosphere,
allowing scientists to reconstruct CO2 levels going back hundreds of
thousands of years and confirm that current concentrations are
unprecedented in that record. Satellite measurements show accelerating
ice loss from the Greenland and Antarctic ice sheets, and tide gauges
combined with satellite altimetry document rising sea levels, driven both
by melting ice and by the thermal expansion of seawater as oceans absorb
excess heat.

The consequences of continued warming are wide-ranging. Rising sea levels
threaten coastal cities and low-lying island nations with increased
flooding and eventual displacement. Warmer oceans fuel more intense
tropical storms and contribute to coral bleaching, which threatens marine
biodiversity. Shifting precipitation patterns are altering agricultural
zones, increasing the frequency and severity of droughts in some regions
while intensifying flooding in others. Heat waves, once rare extreme
events, are becoming more frequent and severe, posing direct risks to
human health, especially among vulnerable populations.

Addressing climate change generally falls into two complementary
categories: mitigation and adaptation. Mitigation refers to efforts to
reduce or prevent greenhouse gas emissions, such as transitioning
electricity generation from coal and natural gas to renewable sources
like solar and wind, electrifying transportation, improving energy
efficiency in buildings and industry, and protecting or restoring
carbon-absorbing ecosystems such as forests and wetlands. Adaptation
refers to adjusting systems and infrastructure to cope with the effects
of climate change that are already locked in, such as building sea walls,
developing drought-resistant crop varieties, and updating building codes
for more extreme weather.

International cooperation has centered around agreements like the 2015
Paris Agreement, in which nearly every country committed to limiting
global warming to well below 2 degrees Celsius above pre-industrial
levels, with an aspirational target of 1.5 degrees. Meeting these targets
requires deep and rapid reductions in global emissions; current national
commitments, even if fully implemented, are widely assessed as
insufficient to meet the 1.5-degree goal, underscoring the scale of the
challenge that remains.

Despite the scale of the problem, the cost of key climate solutions has
fallen dramatically over the past decade. The cost of solar photovoltaic
electricity and battery storage has dropped by well over 80% since 2010,
making renewable energy the cheapest source of new electricity generation
in most parts of the world -- a shift that is reshaping energy investment
decisions independent of climate policy.
""",
}


def write_text_docs():
    for name, content in TEXT_DOCS.items():
        with open(path(name), "w", encoding="utf-8") as f:
            f.write(content.strip() + "\n")
        print(f"  wrote {name} ({os.path.getsize(path(name))} bytes)")


# ---------------------------------------------------------------------------
# 2. PDF documents (4) -- built with reportlab
# ---------------------------------------------------------------------------

PDF_DOCS = {
    "doc06_machine_learning.pdf": (
        "Machine Learning Fundamentals",
        """Machine learning is a subfield of artificial intelligence in which
systems learn patterns from data rather than following explicitly
programmed rules. Instead of a developer hand-coding every decision, a
model is trained on examples and gradually adjusts its internal
parameters to minimize the difference between its predictions and the
correct answers.

There are three broad categories of machine learning. Supervised learning
trains a model on labeled examples, such as pairs of house features and
sale prices, so it can predict labels for new, unseen inputs. Common
supervised algorithms include linear and logistic regression, decision
trees, random forests, gradient-boosted trees, and neural networks.

Unsupervised learning works with unlabeled data, looking for structure or
patterns without predefined correct answers. Clustering algorithms like
k-means group similar data points together, while dimensionality
reduction techniques like principal component analysis (PCA) compress
high-dimensional data into a smaller number of informative dimensions,
often used for visualization or as a preprocessing step.

Reinforcement learning trains an agent to make sequences of decisions by
rewarding desirable outcomes and penalizing undesirable ones, allowing it
to learn optimal strategies through trial and error. This approach has
been notably successful in game-playing systems and robotics.

Deep learning, a subset of machine learning based on artificial neural
networks with many layers, has driven most of the recent breakthroughs in
AI, including large language models. These networks are loosely inspired
by biological neurons: each layer transforms its input through a set of
weighted connections and a nonlinear activation function, allowing the
overall network to approximate extremely complex functions given enough
data and computation.

A model's ability to generalize -- to perform well on new data it has
never seen -- is the central challenge of machine learning. Overfitting
occurs when a model memorizes the training data too closely, including
its noise, and fails to generalize to new examples. Techniques such as
regularization, dropout, cross-validation, and simply gathering more
training data all help control overfitting.

Embeddings are a particularly important machine learning concept for
retrieval systems. An embedding model maps a piece of text, an image, or
another type of data into a dense vector of numbers such that
semantically similar inputs are mapped to nearby points in that vector
space. This property is what makes vector search possible: rather than
matching exact keywords, a retrieval system can find documents whose
meaning is close to a query's meaning, even if they share no words in
common.

Evaluating machine learning systems requires carefully chosen metrics.
Accuracy alone can be misleading, especially for imbalanced datasets;
precision, recall, and the F1 score provide a more nuanced picture of
classification performance, while metrics like mean squared error or
mean absolute error are typically used to evaluate regression models.""",
    ),
    "doc07_ancient_rome.pdf": (
        "The Rise and Fall of Ancient Rome",
        """According to tradition, Rome was founded in 753 BCE by Romulus on the
banks of the Tiber River in central Italy. For its first two and a half
centuries, Rome was ruled by a series of kings, until the last king,
Tarquin the Proud, was overthrown in 509 BCE and the Roman Republic was
established.

The Republic was governed by elected magistrates and two annually elected
consuls, along with the Senate, a body of aristocratic advisors. Over the
following centuries, Rome expanded its territory through a combination of
military conquest and shrewd diplomacy, eventually controlling the entire
Italian peninsula, and later the western Mediterranean following its
victories over Carthage in the three Punic Wars (264-146 BCE). The most
famous episode of these wars was the campaign of the Carthaginian general
Hannibal, who famously led an army, including war elephants, across the
Alps to invade Italy directly.

The late Republic was marked by growing social and political instability,
including a series of civil wars driven by ambitious generals commanding
personally loyal armies. Julius Caesar's rise to power, his crossing of
the Rubicon River in 49 BCE in defiance of the Senate, and his eventual
assassination in 44 BCE by a group of senators who feared he intended to
make himself king, marked the effective end of the Republic. Following
another round of civil war, Caesar's adopted heir Octavian emerged
victorious and, in 27 BCE, was granted the title Augustus, becoming the
first Roman Emperor and beginning the Roman Empire.

The Pax Romana, a roughly 200-year period of relative peace and stability
beginning with Augustus's reign, saw the Empire reach its greatest
territorial extent under Emperor Trajan in the early 2nd century CE,
stretching from Britain in the northwest to Mesopotamia in the east and
Egypt in the south. Roman engineering achievements from this era --
aqueducts, road networks, amphitheaters like the Colosseum, and
sophisticated concrete construction -- still influence infrastructure
design today.

Roman law, embodied in codifications such as the later Corpus Juris
Civilis under Emperor Justinian, established legal principles -- including
concepts of due process, contracts, and property rights -- that remain
foundational to many modern civil law systems. Latin, the language of
Rome, evolved into the Romance languages (Italian, French, Spanish,
Portuguese, and Romanian) and left an enduring imprint on English
vocabulary and Western scientific and legal terminology.

The Western Roman Empire faced mounting pressures from the 3rd century CE
onward: economic instability, political fragmentation, plague, and
increasing incursions by various groups along its northern and eastern
frontiers. In 476 CE, the Germanic leader Odoacer deposed the last
Western Roman Emperor, Romulus Augustulus, an event conventionally marked
as the fall of the Western Roman Empire, though the Eastern Roman Empire,
commonly known as the Byzantine Empire, continued for nearly another
thousand years until the fall of Constantinople in 1453.

The causes of Rome's decline remain debated among historians, but
commonly cited factors include overextension of the Empire's borders,
economic strain from a debased currency and heavy taxation, political
instability caused by frequent changes in leadership, and the increasing
military and administrative burden of defending an enormous, diverse
territory.""",
    ),
    "doc08_deep_sea_creatures.pdf": (
        "Life in the Deep Sea",
        """The deep sea, generally defined as ocean waters below 200 meters where
sunlight can no longer support photosynthesis, is the largest habitat on
Earth by volume, yet it remains one of the least explored. Scientists
estimate that more than 80% of the ocean has never been mapped, explored,
or even observed by humans.

Life in the deep sea has evolved remarkable adaptations to survive in an
environment characterized by crushing pressure, near-freezing
temperatures, and permanent darkness. Many deep-sea organisms produce
their own light through bioluminescence, a chemical reaction typically
involving a light-emitting molecule called luciferin and an enzyme called
luciferase. The anglerfish is perhaps the most famous example, using a
glowing lure to attract prey in the pitch-black water. Bioluminescence
also serves other purposes, including communication, camouflage through a
technique called counter-illumination, and startling predators.

Hydrothermal vents, discovered in 1977 near the Galapagos Islands,
revealed one of the most surprising ecosystems on the planet. These vents
release mineral-rich, superheated water from beneath the ocean floor.
Rather than relying on sunlight and photosynthesis, the base of the food
web around hydrothermal vents is built on chemosynthesis: specialized
bacteria convert chemicals like hydrogen sulfide, which would be toxic to
most life, into usable energy. Giant tube worms, some growing over two
meters long, host these bacteria symbiotically inside their bodies and
have no digestive system of their own, relying entirely on their
bacterial partners for nutrition.

The giant and colossal squid, among the largest invertebrates on Earth,
inhabit the deep ocean and were, for centuries, known primarily from
carcasses that washed ashore or were found in the stomachs of sperm
whales, their primary predator. It was not until 2004 that a giant squid
was first photographed alive in its natural habitat, and video footage of
a live giant squid in the deep sea was not captured until 2012.

The deep sea also hosts some of the oldest known living organisms.
Certain deep-sea corals have been dated to be over 4,000 years old, and
some glass sponges are estimated to live for thousands of years, growing
extraordinarily slowly in the cold, stable conditions of the deep ocean
floor.

Pressure at extreme depths, such as the roughly 11,000-meter-deep
Challenger Deep in the Mariana Trench, exceeds 1,000 times atmospheric
pressure at sea level. Organisms that live at these depths have evolved
flexible cell membranes and specialized proteins that remain functional
under pressures that would crush organisms adapted to shallower waters.

Deep-sea ecosystems face growing threats from human activity, including
proposed deep-sea mining for polymetallic nodules containing cobalt,
nickel, and manganese -- minerals in high demand for battery
manufacturing. Because deep-sea ecosystems are slow-growing and poorly
understood, scientists have raised significant concern that mining
activity could cause irreversible damage before these ecosystems are
even fully cataloged, spurring ongoing international debate over
regulation of the emerging deep-sea mining industry.""",
    ),
    "doc09_renewable_energy.pdf": (
        "Renewable Energy Technologies",
        """Renewable energy refers to power generated from naturally replenishing
sources, in contrast to finite fossil fuels like coal, oil, and natural
gas. The main renewable technologies deployed at scale today are solar,
wind, hydroelectric, geothermal, and bioenergy, each with distinct
technical characteristics, costs, and best-fit geographic conditions.

Solar photovoltaic (PV) panels convert sunlight directly into electricity
using semiconductor materials, typically silicon, that exhibit the
photovoltaic effect: photons striking the material knock electrons loose,
generating a flow of current. Solar PV costs have fallen more than 80%
over the past decade, driven by manufacturing scale, technological
improvements, and supply chain maturation, making it one of the cheapest
sources of new electricity generation in much of the world. Solar
thermal systems, a related but distinct technology, instead use mirrors
to concentrate sunlight to generate heat, which can drive a conventional
steam turbine or be stored for later use, providing a path to dispatchable
solar power even after the sun sets.

Wind turbines convert the kinetic energy of moving air into electricity
using large rotating blades connected to a generator. Modern onshore wind
turbines commonly stand over 100 meters tall with blades spanning 60
meters or more, while offshore wind farms, though more expensive to build
and maintain, can access stronger and more consistent wind resources and
avoid many land-use conflicts.

Hydroelectric power, the oldest and historically largest source of
renewable electricity, generates power by channeling flowing or falling
water through turbines. Large dam projects can provide substantial,
reliable baseload power and flood control but often carry significant
ecological and social costs, including habitat disruption and
displacement of communities. Smaller "run-of-river" hydro projects avoid
large reservoirs and have a comparatively lighter environmental
footprint, though they generate less power and are more sensitive to
seasonal flow variation.

Geothermal energy taps heat from within the Earth, either for direct
heating applications or, at higher temperatures, to generate electricity
by driving steam turbines. Geothermal resources are geographically
concentrated in areas with high volcanic or tectonic activity, such as
Iceland, which derives a substantial share of its electricity and nearly
all of its heating from geothermal sources.

Bioenergy converts organic matter -- crops, agricultural residue, wood,
or organic waste -- into usable energy, either through direct combustion,
conversion into biogas, or processing into liquid biofuels like ethanol
and biodiesel. Bioenergy's climate benefit depends heavily on the source
material and land-use practices; biofuels grown on land converted from
forest, for example, can have a substantially worse climate impact than
the fossil fuels they replace.

A central challenge for solar and wind power is intermittency: neither
generates electricity at a constant, predictable rate, since output
depends on weather and time of day. Addressing this challenge is an
active area of technology development, including grid-scale battery
storage, pumped-hydro storage, improved long-distance transmission to
balance supply across wider geographic areas, and demand-response
programs that shift electricity usage to match availability. As storage
costs continue to fall alongside generation costs, renewable energy is
increasingly able to provide reliable, dispatchable power around the
clock.""",
    ),
}


def write_pdf_docs():
    for name, (title, body) in PDF_DOCS.items():
        filepath = path(name)
        c = canvas.Canvas(filepath, pagesize=LETTER)
        width, height = LETTER
        margin = 0.9 * inch
        usable_width = width - 2 * margin

        c.setFont("Helvetica-Bold", 16)
        y = height - margin
        c.drawString(margin, y, title)
        y -= 0.35 * inch

        c.setFont("Helvetica", 10.5)
        wrap_width_chars = 95
        for paragraph in body.strip().split("\n\n"):
            lines = textwrap.wrap(" ".join(paragraph.split()), wrap_width_chars)
            for line in lines:
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 10.5)
                    y = height - margin
                c.drawString(margin, y, line)
                y -= 14
            y -= 10  # paragraph spacing

        c.save()
        print(f"  wrote {name} ({os.path.getsize(filepath)} bytes)")


# ---------------------------------------------------------------------------
# 3. Image documents (4) -- rendered text "scans", for OCR extraction
# ---------------------------------------------------------------------------

IMAGE_DOCS = {
    "doc10_company_memo.png": (
        "INTERNAL MEMO",
        """From: Operations Team
To: All Staff
Subject: Updated Office Hours Starting Next Month

Starting the first of next month, our main office will shift to a hybrid
schedule. Employees are asked to be present in the office on Tuesdays,
Wednesdays, and Thursdays, with Mondays and Fridays available for remote
work at each team's discretion.

Core collaboration hours remain 10 AM to 3 PM local time on all in-office
days, during which staff should be reachable for meetings. The building
badge system will be updated to reflect the new schedule by the end of
this week.

Please direct any scheduling conflicts or accommodation requests to your
direct manager. Facilities will also be adjusting desk booking capacity
to match the new attendance pattern.

Thank you for your continued flexibility as we roll out this change.""",
    ),
    "doc11_recipe_card.png": (
        "RECIPE: Simple Tomato Basil Soup",
        """Ingredients:
- 2 tablespoons olive oil
- 1 diced yellow onion
- 3 cloves garlic, minced
- 2 cans (28 oz each) whole peeled tomatoes
- 2 cups vegetable stock
- 1/2 cup heavy cream
- A handful of fresh basil leaves
- Salt and pepper to taste

Instructions:
1. Heat olive oil in a large pot over medium heat. Add onion and cook
   until softened, about 5 minutes.
2. Add garlic and cook for another minute until fragrant.
3. Add tomatoes (with juice) and vegetable stock. Bring to a simmer and
   cook for 20 minutes, breaking up the tomatoes with a spoon.
4. Blend the soup until smooth using an immersion blender, or carefully
   transfer to a countertop blender in batches.
5. Stir in the cream and torn basil leaves. Season with salt and pepper.
6. Simmer for another 5 minutes and serve warm with crusty bread.

Serves 4. Total time: about 35 minutes.""",
    ),
    "doc12_meeting_notes.png": (
        "MEETING NOTES - Product Sync",
        """Date: Weekly Product Sync
Attendees: Product, Engineering, Design leads

Agenda item 1 - Q3 roadmap review:
The search relevance project is on track for a mid-quarter release.
Engineering flagged that the new ranking model will require an extra
week of load testing before rollout.

Agenda item 2 - Customer feedback themes:
Support tickets this month were dominated by requests for better export
options and complaints about onboarding being too long. Design will
mock up a shortened onboarding flow for review next week.

Agenda item 3 - Metrics review:
Weekly active users are up 8 percent month over month. Churn on the
starter plan increased slightly, likely tied to the recent pricing
change; finance will pull a cohort analysis before the next meeting.

Action items:
- Engineering: finish load testing plan by Friday
- Design: onboarding mockups by next Wednesday
- Product: draft pricing cohort analysis request""",
    ),
    "doc13_product_spec.png": (
        "PRODUCT SPEC: Notification Preferences v2",
        """Overview:
This spec covers the redesign of user notification preferences, allowing
granular control over channel (email, push, SMS) and category
(billing, security, product updates, marketing).

Requirements:
- Users must be able to toggle each category independently per channel.
- Security notifications cannot be fully disabled, only routed to email.
- Changes should take effect immediately without requiring a page reload.
- Default settings for new users: email on for all categories, push on
  for security and billing only, SMS off by default.

Non-goals:
- This spec does not cover in-app notification center redesign.
- Digest/summary email frequency settings are tracked separately.

Success metrics:
- Reduce notification-related support tickets by 25 percent.
- Increase opt-in rate for product update emails among engaged users.""",
    ),
}


def _load_font(size: int):
    candidates = [
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def write_image_docs():
    title_font = _load_font(20)
    body_font = _load_font(15)

    for name, (title, body) in IMAGE_DOCS.items():
        width, height = 700, 900
        img = Image.new("L", (width, height), color=255)  # grayscale, smaller file
        draw = ImageDraw.Draw(img)

        margin = 35
        y = margin
        draw.text((margin, y), title, fill=0, font=title_font)
        y += 38
        draw.line((margin, y, width - margin, y), fill=0, width=2)
        y += 20

        for paragraph in body.strip().split("\n\n"):
            wrapped = textwrap.wrap(paragraph, width=68)
            for line in wrapped:
                draw.text((margin, y), line, fill=0, font=body_font)
                y += 21
            y += 10

        filepath = path(name)
        img.save(filepath, "PNG", optimize=True)
        print(f"  wrote {name} ({os.path.getsize(filepath)} bytes)")


# ---------------------------------------------------------------------------
# 4. Source-code "documents" (4) -- realistic scripts, indexed as text
# ---------------------------------------------------------------------------

CODE_DOCS = {
    "doc14_data_pipeline.py": '''"""
data_pipeline.py

A small ETL pipeline that extracts customer order records from a CSV
export, cleans and transforms them, and loads the results into a local
SQLite database. Included here as a sample "document" so the text
extraction step can demonstrate reading source code the same way it
reads prose.
"""

import csv
import sqlite3
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

DB_PATH = Path("orders.db")
CSV_PATH = Path("raw_orders.csv")


@dataclass
class Order:
    order_id: str
    customer_email: str
    amount_cents: int
    currency: str
    order_date: datetime
    status: str

    @classmethod
    def from_row(cls, row: dict) -> "Order":
        return cls(
            order_id=row["order_id"].strip(),
            customer_email=row["email"].strip().lower(),
            amount_cents=round(float(row["amount"]) * 100),
            currency=row.get("currency", "USD").upper(),
            order_date=datetime.strptime(row["date"], "%Y-%m-%d"),
            status=row.get("status", "unknown").lower(),
        )


def extract(csv_path: Path) -> list[dict]:
    """Read raw rows from the CSV export."""
    logger.info("Extracting rows from %s", csv_path)
    with csv_path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    logger.info("Extracted %d raw rows", len(rows))
    return rows


def transform(rows: list[dict]) -> list[Order]:
    """Clean and validate raw rows, dropping any that fail parsing."""
    orders: list[Order] = []
    skipped = 0
    for row in rows:
        try:
            order = Order.from_row(row)
        except (KeyError, ValueError) as exc:
            logger.warning("Skipping malformed row: %s (%s)", row, exc)
            skipped += 1
            continue

        if order.amount_cents <= 0:
            logger.warning("Skipping non-positive amount for order %s", order.order_id)
            skipped += 1
            continue

        orders.append(order)

    logger.info("Transformed %d orders, skipped %d", len(orders), skipped)
    return orders


def load(orders: list[Order], db_path: Path) -> None:
    """Write cleaned orders into a SQLite table, replacing any existing data."""
    conn = sqlite3.connect(db_path)
    try:
        conn.execute("DROP TABLE IF EXISTS orders")
        conn.execute(
            """
            CREATE TABLE orders (
                order_id TEXT PRIMARY KEY,
                customer_email TEXT NOT NULL,
                amount_cents INTEGER NOT NULL,
                currency TEXT NOT NULL,
                order_date TEXT NOT NULL,
                status TEXT NOT NULL
            )
            """
        )
        conn.executemany(
            """
            INSERT INTO orders (order_id, customer_email, amount_cents, currency, order_date, status)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            [
                (
                    o.order_id,
                    o.customer_email,
                    o.amount_cents,
                    o.currency,
                    o.order_date.isoformat(),
                    o.status,
                )
                for o in orders
            ],
        )
        conn.commit()
        logger.info("Loaded %d orders into %s", len(orders), db_path)
    finally:
        conn.close()


def summarize(orders: list[Order]) -> dict:
    """Produce a quick summary used for a post-run sanity check."""
    total_cents = sum(o.amount_cents for o in orders)
    by_status: dict[str, int] = {}
    for o in orders:
        by_status[o.status] = by_status.get(o.status, 0) + 1

    return {
        "order_count": len(orders),
        "total_amount": total_cents / 100,
        "status_breakdown": by_status,
    }


def run_pipeline(csv_path: Path = CSV_PATH, db_path: Path = DB_PATH) -> dict:
    rows = extract(csv_path)
    orders = transform(rows)
    load(orders, db_path)
    summary = summarize(orders)
    logger.info("Pipeline summary: %s", summary)
    return summary


if __name__ == "__main__":
    run_pipeline()
''',
    "doc15_api_server.js": '''/**
 * api_server.js
 *
 * A minimal Express.js REST API for managing a to-do list, backed by an
 * in-memory store. Included as a sample source-code "document" so the
 * text-extraction step can be demonstrated on JavaScript files too.
 */

const express = require("express");
const { randomUUID } = require("crypto");

const app = express();
app.use(express.json());

const PORT = process.env.PORT || 3000;

/** @type {Map<string, {id: string, title: string, done: boolean, createdAt: string}>} */
const todos = new Map();

function validateTitle(title) {
  if (typeof title !== "string") return "title must be a string";
  const trimmed = title.trim();
  if (trimmed.length === 0) return "title must not be empty";
  if (trimmed.length > 200) return "title must be 200 characters or fewer";
  return null;
}

// Create a new todo item
app.post("/todos", (req, res) => {
  const { title } = req.body ?? {};
  const error = validateTitle(title);
  if (error) {
    return res.status(400).json({ error });
  }

  const todo = {
    id: randomUUID(),
    title: title.trim(),
    done: false,
    createdAt: new Date().toISOString(),
  };
  todos.set(todo.id, todo);
  return res.status(201).json(todo);
});

// List all todos, optionally filtered by completion status
app.get("/todos", (req, res) => {
  let items = Array.from(todos.values());

  if (req.query.done === "true") {
    items = items.filter((t) => t.done);
  } else if (req.query.done === "false") {
    items = items.filter((t) => !t.done);
  }

  items.sort((a, b) => a.createdAt.localeCompare(b.createdAt));
  return res.json(items);
});

// Fetch a single todo by id
app.get("/todos/:id", (req, res) => {
  const todo = todos.get(req.params.id);
  if (!todo) {
    return res.status(404).json({ error: "todo not found" });
  }
  return res.json(todo);
});

// Update a todo's title and/or done status
app.patch("/todos/:id", (req, res) => {
  const todo = todos.get(req.params.id);
  if (!todo) {
    return res.status(404).json({ error: "todo not found" });
  }

  const { title, done } = req.body ?? {};

  if (title !== undefined) {
    const error = validateTitle(title);
    if (error) {
      return res.status(400).json({ error });
    }
    todo.title = title.trim();
  }

  if (done !== undefined) {
    if (typeof done !== "boolean") {
      return res.status(400).json({ error: "done must be a boolean" });
    }
    todo.done = done;
  }

  todos.set(todo.id, todo);
  return res.json(todo);
});

// Delete a todo
app.delete("/todos/:id", (req, res) => {
  const existed = todos.delete(req.params.id);
  if (!existed) {
    return res.status(404).json({ error: "todo not found" });
  }
  return res.status(204).send();
});

// Simple health check endpoint, useful for load balancers / uptime checks
app.get("/healthz", (_req, res) => {
  res.json({ status: "ok", todoCount: todos.size });
});

app.use((err, _req, res, _next) => {
  console.error("Unhandled error:", err);
  res.status(500).json({ error: "internal server error" });
});

app.listen(PORT, () => {
  console.log(`Todo API listening on port ${PORT}`);
});

module.exports = app;
''',
    "doc16_database_schema.sql": """-- database_schema.sql
--
-- Schema for a small e-commerce application: customers, products,
-- orders, and order line items. Included as a sample source-code
-- "document" to demonstrate text extraction on SQL files.

CREATE TABLE customers (
    customer_id     SERIAL PRIMARY KEY,
    email           VARCHAR(255) NOT NULL UNIQUE,
    full_name       VARCHAR(255) NOT NULL,
    hashed_password VARCHAR(255) NOT NULL,
    created_at      TIMESTAMP NOT NULL DEFAULT now(),
    is_active       BOOLEAN NOT NULL DEFAULT TRUE
);

CREATE TABLE addresses (
    address_id   SERIAL PRIMARY KEY,
    customer_id  INTEGER NOT NULL REFERENCES customers(customer_id) ON DELETE CASCADE,
    line1        VARCHAR(255) NOT NULL,
    line2        VARCHAR(255),
    city         VARCHAR(100) NOT NULL,
    state        VARCHAR(100),
    postal_code  VARCHAR(20) NOT NULL,
    country      VARCHAR(2) NOT NULL,
    is_default   BOOLEAN NOT NULL DEFAULT FALSE
);

CREATE TABLE categories (
    category_id   SERIAL PRIMARY KEY,
    name          VARCHAR(100) NOT NULL UNIQUE,
    parent_id     INTEGER REFERENCES categories(category_id)
);

CREATE TABLE products (
    product_id    SERIAL PRIMARY KEY,
    sku           VARCHAR(64) NOT NULL UNIQUE,
    name          VARCHAR(255) NOT NULL,
    description   TEXT,
    category_id   INTEGER REFERENCES categories(category_id),
    price_cents   INTEGER NOT NULL CHECK (price_cents >= 0),
    currency      CHAR(3) NOT NULL DEFAULT 'USD',
    stock_qty     INTEGER NOT NULL DEFAULT 0 CHECK (stock_qty >= 0),
    created_at    TIMESTAMP NOT NULL DEFAULT now()
);

CREATE TABLE orders (
    order_id       SERIAL PRIMARY KEY,
    customer_id    INTEGER NOT NULL REFERENCES customers(customer_id),
    shipping_address_id INTEGER NOT NULL REFERENCES addresses(address_id),
    status         VARCHAR(20) NOT NULL DEFAULT 'pending'
                   CHECK (status IN ('pending', 'paid', 'shipped', 'delivered', 'cancelled')),
    placed_at      TIMESTAMP NOT NULL DEFAULT now(),
    total_cents    INTEGER NOT NULL CHECK (total_cents >= 0)
);

CREATE TABLE order_items (
    order_item_id  SERIAL PRIMARY KEY,
    order_id       INTEGER NOT NULL REFERENCES orders(order_id) ON DELETE CASCADE,
    product_id     INTEGER NOT NULL REFERENCES products(product_id),
    quantity       INTEGER NOT NULL CHECK (quantity > 0),
    unit_price_cents INTEGER NOT NULL CHECK (unit_price_cents >= 0)
);

-- Helpful indexes for common query patterns
CREATE INDEX idx_orders_customer_id ON orders(customer_id);
CREATE INDEX idx_orders_status ON orders(status);
CREATE INDEX idx_order_items_order_id ON order_items(order_id);
CREATE INDEX idx_products_category_id ON products(category_id);
CREATE INDEX idx_addresses_customer_id ON addresses(customer_id);

-- View: order totals reconciled against their line items, useful for
-- catching bugs where the stored total drifts from the computed total.
CREATE VIEW order_total_check AS
SELECT
    o.order_id,
    o.total_cents AS stored_total_cents,
    SUM(oi.quantity * oi.unit_price_cents) AS computed_total_cents,
    o.total_cents - SUM(oi.quantity * oi.unit_price_cents) AS difference_cents
FROM orders o
JOIN order_items oi ON oi.order_id = o.order_id
GROUP BY o.order_id, o.total_cents;

-- View: best-selling products by total quantity sold
CREATE VIEW best_selling_products AS
SELECT
    p.product_id,
    p.name,
    SUM(oi.quantity) AS total_units_sold,
    SUM(oi.quantity * oi.unit_price_cents) AS total_revenue_cents
FROM products p
JOIN order_items oi ON oi.product_id = p.product_id
JOIN orders o ON o.order_id = oi.order_id
WHERE o.status IN ('paid', 'shipped', 'delivered')
GROUP BY p.product_id, p.name
ORDER BY total_units_sold DESC;
""",
    "doc17_deploy.sh": """#!/usr/bin/env bash
#
# deploy.sh
#
# Deploys the web application to a target environment (staging or
# production). Included as a sample source-code "document" to
# demonstrate text extraction on shell scripts.

set -euo pipefail

ENVIRONMENT="${1:-staging}"
APP_NAME="rag-demo-app"
DOCKER_REGISTRY="registry.example.com/${APP_NAME}"
HEALTHCHECK_URL_STAGING="https://staging.example.com/healthz"
HEALTHCHECK_URL_PROD="https://app.example.com/healthz"

log() {
    echo "[deploy] $(date -u +'%Y-%m-%dT%H:%M:%SZ') $*"
}

fail() {
    echo "[deploy] ERROR: $*" >&2
    exit 1
}

if [[ "${ENVIRONMENT}" != "staging" && "${ENVIRONMENT}" != "production" ]]; then
    fail "Unknown environment '${ENVIRONMENT}'. Use 'staging' or 'production'."
fi

log "Starting deployment to ${ENVIRONMENT}"

# 1. Run the test suite before doing anything else.
log "Running test suite..."
if ! npm test --silent; then
    fail "Test suite failed, aborting deployment."
fi

# 2. Build and tag the Docker image.
GIT_SHA="$(git rev-parse --short HEAD)"
IMAGE_TAG="${DOCKER_REGISTRY}:${ENVIRONMENT}-${GIT_SHA}"
log "Building Docker image ${IMAGE_TAG}"
docker build -t "${IMAGE_TAG}" .

# 3. Push the image to the registry.
log "Pushing image to registry..."
docker push "${IMAGE_TAG}"

# 4. Trigger the deployment via the orchestration platform's CLI.
log "Deploying ${IMAGE_TAG} to ${ENVIRONMENT}..."
case "${ENVIRONMENT}" in
    staging)
        kubectl --context staging set image deployment/"${APP_NAME}" "${APP_NAME}=${IMAGE_TAG}"
        HEALTHCHECK_URL="${HEALTHCHECK_URL_STAGING}"
        ;;
    production)
        read -r -p "Confirm production deployment of ${IMAGE_TAG}? [y/N] " CONFIRM
        if [[ "${CONFIRM}" != "y" && "${CONFIRM}" != "Y" ]]; then
            fail "Production deployment cancelled by user."
        fi
        kubectl --context production set image deployment/"${APP_NAME}" "${APP_NAME}=${IMAGE_TAG}"
        HEALTHCHECK_URL="${HEALTHCHECK_URL_PROD}"
        ;;
esac

# 5. Wait for the rollout to complete.
KUBE_CONTEXT="${ENVIRONMENT}"
log "Waiting for rollout to complete..."
kubectl --context "${KUBE_CONTEXT}" rollout status deployment/"${APP_NAME}" --timeout=180s \\
    || fail "Rollout did not complete in time."

# 6. Verify the health check endpoint responds successfully.
log "Verifying health check at ${HEALTHCHECK_URL}"
ATTEMPTS=0
MAX_ATTEMPTS=10
until curl --fail --silent --output /dev/null "${HEALTHCHECK_URL}"; do
    ATTEMPTS=$((ATTEMPTS + 1))
    if [[ "${ATTEMPTS}" -ge "${MAX_ATTEMPTS}" ]]; then
        fail "Health check did not pass after ${MAX_ATTEMPTS} attempts."
    fi
    log "Health check attempt ${ATTEMPTS} failed, retrying in 5s..."
    sleep 5
done

log "Deployment to ${ENVIRONMENT} completed successfully (image: ${IMAGE_TAG})"
""",
}


def write_code_docs():
    for name, content in CODE_DOCS.items():
        filepath = path(name)
        with open(filepath, "w", encoding="utf-8", newline="\n") as f:
            f.write(content)
        print(f"  wrote {name} ({os.path.getsize(filepath)} bytes)")


# ---------------------------------------------------------------------------
# 5. Word documents (3) -- built with python-docx
# ---------------------------------------------------------------------------

DOCX_DOCS = {
    "doc18_project_proposal.docx": (
        "Project Proposal: Internal Knowledge Search Tool",
        [
            ("heading", "Background"),
            ("para", "Employees across support, sales, and engineering currently "
                      "spend a significant amount of time searching for information "
                      "scattered across wikis, shared drives, PDFs, and old chat "
                      "threads. Informal surveys suggest staff spend an average of "
                      "close to five hours per week simply locating documents they "
                      "know exist somewhere but cannot easily find."),
            ("heading", "Proposed Solution"),
            ("para", "We propose building an internal knowledge search tool based on "
                      "a Retrieval-Augmented Generation (RAG) architecture. Source "
                      "documents -- including text files, PDFs, scanned images, source "
                      "code, and Word documents -- would be ingested, split into "
                      "chunks, embedded into a vector representation, and indexed in "
                      "a vector database. When an employee asks a question, the "
                      "system retrieves the most relevant chunks and passes them, "
                      "along with the question, to a large language model to "
                      "generate a grounded, cited answer."),
            ("heading", "Scope"),
            ("para", "Phase 1 will cover ingestion of the top five most-referenced "
                      "internal document sources, a basic web search interface, and "
                      "citation links back to the original source document. Phase 2 "
                      "will expand ingestion to additional sources and add support "
                      "for access-control-aware retrieval, so that search results "
                      "respect existing document permissions."),
            ("heading", "Success Metrics"),
            ("para", "We will consider Phase 1 successful if at least 60 percent of "
                      "pilot users report the tool saved them time in a follow-up "
                      "survey, and if median time-to-answer for common support "
                      "questions drops by at least 30 percent compared to the "
                      "current manual search process."),
            ("heading", "Timeline and Resources"),
            ("para", "We estimate Phase 1 will take approximately eight weeks with "
                      "two engineers and one part-time designer. Phase 2 scope and "
                      "timeline will be finalized based on Phase 1 pilot feedback."),
            ("heading", "Risks"),
            ("para", "Key risks include the quality of retrieval on poorly "
                      "structured legacy documents, the cost of running embedding "
                      "and generation calls at scale, and the need to keep the "
                      "index synchronized as source documents are updated or "
                      "deleted. We plan to mitigate these with a scheduled "
                      "re-indexing job and a smaller pilot document set before "
                      "expanding coverage."),
        ],
    ),
    "doc19_employee_handbook.docx": (
        "Employee Handbook Excerpt: Time Off Policy",
        [
            ("heading", "Overview"),
            ("para", "This section describes the company's paid time off (PTO) "
                      "policy, which applies to all full-time employees beginning "
                      "on their first day of employment."),
            ("heading", "Accrual"),
            ("para", "Full-time employees accrue PTO at a rate of 1.67 days per "
                      "month, totaling 20 days per calendar year. PTO accrual "
                      "begins immediately upon hire, though new employees are "
                      "asked to avoid scheduling extended time off during their "
                      "first 90 days unless discussed with their manager in "
                      "advance."),
            ("heading", "Requesting Time Off"),
            ("para", "Employees should submit PTO requests through the HR portal "
                      "at least two weeks in advance for planned absences longer "
                      "than three consecutive days. Shorter absences and sick days "
                      "may be requested with less notice. Managers are expected to "
                      "respond to PTO requests within three business days."),
            ("heading", "Carryover and Payout"),
            ("para", "Employees may carry over up to five unused PTO days into the "
                      "following calendar year; any balance beyond that is "
                      "forfeited unless required otherwise by local law. Upon "
                      "voluntary or involuntary termination, unused accrued PTO "
                      "will be paid out in accordance with applicable state or "
                      "national regulations."),
            ("heading", "Holidays"),
            ("para", "In addition to PTO, the company observes ten paid holidays "
                      "per year, published annually on the HR portal. Employees in "
                      "regions with different statutory holidays should refer to "
                      "their region-specific holiday calendar."),
            ("heading", "Sick Leave"),
            ("para", "Sick leave is tracked separately from general PTO in "
                      "jurisdictions where required by law. Employees who are "
                      "unwell are encouraged to stay home and use sick leave rather "
                      "than PTO, and should notify their manager as early as "
                      "possible on the day of absence."),
        ],
    ),
    "doc20_travel_guide.docx": (
        "Travel Guide: Visiting Kyoto, Japan",
        [
            ("heading", "Best Time to Visit"),
            ("para", "Kyoto is a year-round destination, but the two most popular "
                      "seasons are spring (late March through April), when cherry "
                      "blossoms bloom across the city's parks and temple grounds, "
                      "and autumn (mid-November), when maple leaves turn brilliant "
                      "shades of red and orange. Both seasons draw large crowds, so "
                      "booking accommodations well in advance is strongly "
                      "recommended."),
            ("heading", "Getting Around"),
            ("para", "Kyoto's historic districts are relatively compact and highly "
                      "walkable, but the city also has an extensive and reliable "
                      "bus network that reaches most major temples and shrines. A "
                      "one-day or two-day bus pass is often more economical than "
                      "paying individual fares if visiting multiple sites in a "
                      "single day. Renting a bicycle is another popular option for "
                      "covering more ground at a relaxed pace."),
            ("heading", "Must-See Sites"),
            ("para", "Fushimi Inari Taisha, famous for its thousands of vermillion "
                      "torii gates winding up the mountainside, is free to enter "
                      "and open 24 hours, making early morning or evening visits a "
                      "good way to avoid the busiest midday crowds. Kinkaku-ji, the "
                      "Golden Pavilion, is a Zen Buddhist temple covered in gold "
                      "leaf that reflects strikingly in its surrounding pond. The "
                      "Arashiyama Bamboo Grove, on the city's western edge, offers "
                      "a serene walk through towering bamboo stalks and is often "
                      "combined with a visit to the nearby Tenryu-ji temple."),
            ("heading", "Food and Dining"),
            ("para", "Kyoto is renowned for kaiseki, a traditional multi-course "
                      "meal emphasizing seasonal ingredients and refined "
                      "presentation. For a more casual experience, Nishiki Market, "
                      "often called 'Kyoto's Kitchen,' is a narrow shopping street "
                      "lined with food stalls selling local specialties, from "
                      "fresh yuba (tofu skin) to skewered street food."),
            ("heading", "Practical Tips"),
            ("para", "Many temples charge a small entry fee, typically between 300 "
                      "and 600 yen, and some close as early as 4:30 or 5 PM, so it "
                      "is worth checking hours in advance. Carrying cash is still "
                      "advisable, as some smaller shops, temples, and restaurants "
                      "do not accept credit cards."),
        ],
    ),
}


def write_docx_docs():
    for name, (title, sections) in DOCX_DOCS.items():
        doc = Document()
        doc.add_heading(title, level=0)
        for kind, text in sections:
            if kind == "heading":
                doc.add_heading(text, level=1)
            else:
                doc.add_paragraph(text)
        filepath = path(name)
        doc.save(filepath)
        print(f"  wrote {name} ({os.path.getsize(filepath)} bytes)")


# ---------------------------------------------------------------------------

def main():
    print("Writing text documents...")
    write_text_docs()
    print("Writing PDF documents...")
    write_pdf_docs()
    print("Writing image documents...")
    write_image_docs()
    print("Writing code documents...")
    write_code_docs()
    print("Writing Word documents...")
    write_docx_docs()

    all_files = sorted(os.listdir(DOCS_DIR))
    print(f"\nDone. {len(all_files)} files written to {DOCS_DIR}:")
    for name in all_files:
        size_kb = os.path.getsize(path(name)) / 1024
        print(f"  {name:<32} {size_kb:6.1f} KB")


if __name__ == "__main__":
    main()
