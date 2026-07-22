"""
01_text_extraction.py

Author: Ram Kumar Lanke
Applied AI Solutions Architect
lankeramkumar@gmail.com

RAG STEP 1: TEXT EXTRACTION
============================
Real-world documents arrive in many formats -- plain text, PDF, Word,
scanned images, source code. Before you can chunk, embed, or retrieve
anything, you need a single, uniform representation: plain text plus a
bit of metadata about where it came from.

This script walks documents/, dispatches each file to a format-specific
extractor, and writes a single unified JSON "corpus" file that every
later step in the pipeline (chunking, embedding, retrieval...) reads
from. That's the general shape of the extraction stage in any RAG
system: many input formats in, one normalized representation out.

Run:
    python scripts/01_text_extraction.py
"""

import json
import os
from shutil import which as shutil_which

from pypdf import PdfReader
from docx import Document
from PIL import Image

try:
    import pytesseract

    _PYTESSERACT_AVAILABLE = True
    # pytesseract only wraps the `tesseract` binary -- it doesn't ship it.
    # If it's not on PATH, point at the common Windows install location
    # (e.g. the UB-Mannheim installer) rather than failing outright.
    _DEFAULT_WIN_TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if not shutil_which("tesseract") and os.path.exists(_DEFAULT_WIN_TESSERACT):
        pytesseract.pytesseract.tesseract_cmd = _DEFAULT_WIN_TESSERACT
except ImportError:
    _PYTESSERACT_AVAILABLE = False

HERE = os.path.dirname(os.path.abspath(__file__))
# Sunrise Medical Center hospital documents (PDF/Word/TXT/images/audio/video),
# organized into department subfolders -- unlike the flat, generic
# documents/ folder 00_generate_sample_documents.py writes to, this one is
# walked recursively (see run_extraction() below).
DOCS_DIR = os.path.join(HERE, "..", "Sunrise_Medical_Center_Documents")
OUTPUT_DIR = os.path.join(HERE, "..", "output")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "01_extracted_corpus.json")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# File extensions that are already plain text, including source code --
# for extraction purposes, a .py file is no different from a .txt file.
PLAIN_TEXT_EXTENSIONS = {".txt", ".py", ".js", ".sql", ".sh", ".md", ".json", ".csv"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a"}
VIDEO_EXTENSIONS = {".mp4", ".mov"}

# Lazily-loaded, module-level so the (~75MB) model is only loaded once per
# run, not once per audio/video file.
_WHISPER_MODEL = None
_WHISPER_MODEL_NAME = "tiny.en"


def _get_whisper_model():
    global _WHISPER_MODEL
    if _WHISPER_MODEL is None:
        from faster_whisper import WhisperModel

        print(f"  (loading faster-whisper '{_WHISPER_MODEL_NAME}' model for audio/video transcription...)")
        _WHISPER_MODEL = WhisperModel(_WHISPER_MODEL_NAME, device="cpu", compute_type="int8")
    return _WHISPER_MODEL


def extract_plain_text(filepath: str) -> str:
    """.txt and source-code files: just read them. No parsing needed."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_pdf(filepath: str) -> str:
    """PDF: pull text page by page. pypdf handles the container format;
    it does NOT do OCR, so an image-only (scanned) PDF would come back
    empty -- that's a case where you'd need to rasterize pages and OCR
    them, the same way we handle .png files below."""
    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)
    return "\n".join(pages)


def extract_docx(filepath: str) -> str:
    """.docx: python-docx exposes the document as a list of paragraph
    objects (headings included as paragraphs with a "Heading" style).
    We also walk any tables, since table cell text is not part of
    doc.paragraphs."""
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_image(filepath: str) -> str:
    """Images ("scanned" documents here): run OCR with pytesseract, which
    wraps the Tesseract OCR engine. This requires the Tesseract binary to
    be installed separately from the Python package -- pytesseract is
    just a thin wrapper around the `tesseract` executable.

    On Windows: install from https://github.com/UB-Mannheim/tesseract/wiki
    and either add it to PATH or set:
        pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
    On macOS: brew install tesseract
    On Debian/Ubuntu: apt-get install tesseract-ocr
    """
    if not _PYTESSERACT_AVAILABLE:
        return "[OCR skipped: pytesseract package not installed]"

    img = Image.open(filepath)
    try:
        text = pytesseract.image_to_string(img)
        return text
    except Exception as exc:  # pytesseract.TesseractNotFoundError, etc.
        return f"[OCR skipped: Tesseract engine not available ({exc})]"


def extract_audio_or_video(filepath: str) -> str:
    """Audio (.mp3/.wav/.m4a) and video (.mp4/.mov): transcribe the spoken
    narration with a local Whisper model via faster-whisper. Unlike the
    other extractors, this doesn't "parse" an existing text layer -- it
    generates text from audio that never had any, via automatic speech
    recognition (ASR). faster-whisper decodes the audio track straight out
    of a video container itself (via PyAV), so no separate ffmpeg
    extraction step is needed here even for .mp4 input.

    This is deliberately the *last* extractor added to this script: it's
    the one format-adapter pattern doesn't make free -- OCR and ASR both
    require a real model, not just a parsing library, which is why they're
    each gated behind an availability check rather than assumed to work.
    """
    try:
        model = _get_whisper_model()
    except ImportError:
        return "[Transcription skipped: faster-whisper package not installed]"

    try:
        segments, _info = model.transcribe(filepath, beam_size=5)
        return " ".join(seg.text.strip() for seg in segments)
    except Exception as exc:
        return f"[Transcription skipped: {exc}]"


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
}


def extract_document(filepath: str, folder: str = "") -> dict:
    """Dispatch a single file to the right extractor and return a
    normalized record. This is the "adapter" pattern: whatever the file
    type, callers downstream only ever see {doc_id, text, metadata}."""
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filename)[1].lower()

    if ext in PLAIN_TEXT_EXTENSIONS:
        text = extract_plain_text(filepath)
        source_type = "code" if ext in {".py", ".js", ".sql", ".sh"} else "text"
    elif ext in IMAGE_EXTENSIONS:
        text = extract_image(filepath)
        source_type = "image_ocr"
    elif ext in AUDIO_EXTENSIONS:
        text = extract_audio_or_video(filepath)
        source_type = "audio_transcript"
    elif ext in VIDEO_EXTENSIONS:
        text = extract_audio_or_video(filepath)
        source_type = "video_transcript"
    elif ext in EXTRACTORS:
        text = EXTRACTORS[ext](filepath)
        source_type = {"pdf": "pdf", "docx": "word"}[ext.lstrip(".")]
    else:
        raise ValueError(f"No extractor registered for extension '{ext}'")

    # Normalize whitespace: collapse runs of blank lines, strip trailing
    # spaces. Extraction libraries are inconsistent about this, and messy
    # whitespace makes later chunking (which often splits on blank lines)
    # unreliable.
    cleaned = "\n".join(line.rstrip() for line in text.splitlines())
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    cleaned = cleaned.strip()

    failure_markers = ("[OCR skipped", "[Transcription skipped")

    return {
        "doc_id": os.path.splitext(filename)[0],
        "filename": filename,
        "folder": folder,
        "source_type": source_type,
        "extension": ext,
        "text": cleaned,
        "char_count": len(cleaned),
        "extraction_ok": len(cleaned) > 0 and not cleaned.startswith(failure_markers),
    }


def run_extraction() -> list[dict]:
    # Walk recursively: the Sunrise Medical Center documents are organized
    # into department subfolders (01_Employee_Handbooks, 02_Emergency_
    # Protocols, ...), unlike the flat documents/ folder this script
    # originally targeted. `folder` is kept as metadata -- useful later for
    # filtering retrieval by department, or just for display.
    entries = []
    for root, _dirs, files in os.walk(DOCS_DIR):
        rel_folder = os.path.relpath(root, DOCS_DIR)
        rel_folder = "" if rel_folder == "." else rel_folder
        for name in files:
            entries.append((os.path.join(root, name), rel_folder))
    entries.sort(key=lambda e: (e[1], os.path.basename(e[0])))

    records = []
    for filepath, folder in entries:
        try:
            record = extract_document(filepath, folder=folder)
        except Exception as exc:
            record = {
                "doc_id": os.path.splitext(os.path.basename(filepath))[0],
                "filename": os.path.basename(filepath),
                "folder": folder,
                "source_type": "unknown",
                "extension": os.path.splitext(filepath)[1].lower(),
                "text": "",
                "char_count": 0,
                "extraction_ok": False,
                "error": str(exc),
            }
        records.append(record)

        status = "OK" if record["extraction_ok"] else "EMPTY/FAILED"
        print(f"  [{record['source_type']:<16}] {record['filename']:<48} "
              f"{record['char_count']:>6} chars  {status}")

    return records


def main():
    print(f"Extracting text from documents in {DOCS_DIR}\n")
    records = run_extraction()

    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(records, f, indent=2, ensure_ascii=False)

    ok_count = sum(1 for r in records if r["extraction_ok"])
    print(f"\nExtracted {ok_count}/{len(records)} documents successfully.")
    print(f"Corpus written to {OUTPUT_PATH}")

    if ok_count < len(records):
        print(
            "\nNote: some documents extracted no text. For .png files, "
            "install the Tesseract OCR engine (see extract_image()'s "
            "docstring). For .mp3/.mp4 files, `pip install faster-whisper` "
            "(see extract_audio_or_video()'s docstring). Then re-run."
        )


if __name__ == "__main__":
    main()
